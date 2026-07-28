from __future__ import annotations

import ast
import csv
import io
import re
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

from app.ingestion.processors.normalizer import remove_repeated_page_furniture
from app.ingestion.structure import BlockKind, ExtractedDocument, StructuralBlock

QUESTION = re.compile(r"^\s*(?:(?:question|q)\s*)?(\d+[A-Za-z]?)\s*[.):\-]\s+\S", re.I)
HEADING_VALUE = re.compile(r"^\s*([A-Za-z][\w &/()\-]{1,100}):(?:\s+\S.*)?$")
CODE_LANGUAGES = {
    ".py": "python",
    ".js": "javascript",
    ".ts": "typescript",
    ".java": "java",
    ".cpp": "cpp",
}


def extract_document(extension: str, data: bytes, *, filename: str = "") -> ExtractedDocument:
    extension = extension.lower()
    if extension == ".pdf":
        return _pdf(data)
    if extension == ".docx":
        return _docx(data)
    if extension in {".html", ".htm"}:
        return _html(data)
    if extension == ".csv":
        return _csv(data, filename)
    if extension in CODE_LANGUAGES:
        return _source(data, extension, filename)
    return _textual(data, markdown=extension == ".md")


def _pdf(data: bytes) -> ExtractedDocument:
    reader = PdfReader(io.BytesIO(data))
    blocks: list[StructuralBlock] = []
    warnings: list[str] = []
    pages: list[str] = []
    for page in reader.pages:
        # Layout mode is safer for columns when supported; fall back for older/problem PDFs.
        try:
            text = page.extract_text(extraction_mode="layout") or ""
        except (KeyError, TypeError, ValueError):
            try:
                text = page.extract_text() or ""
            except KeyError:
                # A valid image-only/blank page may have no /Contents stream.
                text = ""
        pages.append(text)
    cleaned_pages = remove_repeated_page_furniture(pages)
    empty = sum(not page.strip() for page in cleaned_pages)
    for page_number, text in enumerate(cleaned_pages, 1):
        if not text.strip():
            continue
        blocks.extend(_lines_to_blocks(text, page=page_number))
    page_count = len(reader.pages)
    scanned = empty / page_count if page_count else 1.0
    if any(len(line) > 220 for block in blocks for line in block.text.splitlines()):
        warnings.append("possible_column_or_reading_order_issue")
    return ExtractedDocument(
        blocks,
        page_count=page_count,
        empty_pages=empty,
        scanned_likelihood=scanned,
        reading_order_warnings=warnings,
        metadata={"format": "pdf"},
    )


def _docx(data: bytes) -> ExtractedDocument:
    document = Document(io.BytesIO(data))
    blocks: list[StructuralBlock] = []
    table_number = 0
    # iter_inner_content preserves paragraph/table source order in python-docx 1.1+.
    items = (
        document.iter_inner_content()
        if hasattr(document, "iter_inner_content")
        else document.paragraphs
    )
    for item in items:
        if hasattr(item, "rows"):
            table_number += 1
            rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in item.rows]
            blocks.append(
                StructuralBlock(
                    "\n".join(rows),
                    BlockKind.TABLE,
                    table_identifier=f"table-{table_number}",
                    row_range=f"1-{len(rows)}",
                )
            )
            continue
        text = item.text.strip()
        if not text:
            continue
        style = (item.style.name if item.style else "").lower()
        if style.startswith("heading"):
            level_match = re.search(r"\d+", style)
            level = int(level_match.group()) if level_match else 1
            blocks.append(
                StructuralBlock(
                    text, BlockKind.HEADING, heading=text, heading_level=level, section=text
                )
            )
        elif "list" in style:
            blocks.append(StructuralBlock(text, BlockKind.LIST_ITEM))
        elif "caption" in style:
            blocks.append(StructuralBlock(text, BlockKind.CAPTION))
        else:
            blocks.extend(_lines_to_blocks(text))
    # Compatibility fallback for older python-docx where tables were not interleaved.
    if not hasattr(document, "iter_inner_content"):
        for table in document.tables:
            table_number += 1
            rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            blocks.append(
                StructuralBlock(
                    "\n".join(rows), BlockKind.TABLE, table_identifier=f"table-{table_number}"
                )
            )
    return ExtractedDocument(blocks, metadata={"format": "docx"})


def _html(data: bytes) -> ExtractedDocument:
    soup = BeautifulSoup(data.decode("utf-8", errors="replace"), "html.parser")
    for tag in soup(
        ["script", "style", "noscript", "nav", "template", "iframe", "object", "embed"]
    ):
        tag.decompose()
    for tag in soup.select(
        "[hidden], [aria-hidden='true'], [style*='display:none'], [style*='display: none']"
    ):
        tag.decompose()
    root = soup.find("article") or soup.find("main") or soup.body or soup
    blocks: list[StructuralBlock] = []
    title = soup.title.get_text(" ", strip=True) if soup.title else ""
    if title:
        blocks.append(StructuralBlock(title, BlockKind.TITLE, heading=title, heading_level=1))
    table_number = 0
    for tag in root.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "table"]):
        if tag.find_parent(["table"]) and tag.name != "table":
            continue
        text = tag.get_text(" ", strip=True)
        if not text:
            continue
        if tag.name.startswith("h"):
            blocks.append(
                StructuralBlock(
                    text,
                    BlockKind.HEADING,
                    heading=text,
                    heading_level=int(tag.name[1]),
                    section=text,
                )
            )
        elif tag.name == "li":
            blocks.append(StructuralBlock(text, BlockKind.LIST_ITEM))
        elif tag.name == "table":
            table_number += 1
            rows = [
                " | ".join(cell.get_text(" ", strip=True) for cell in row.find_all(["th", "td"]))
                for row in tag.find_all("tr")
            ]
            blocks.append(
                StructuralBlock(
                    "\n".join(filter(None, rows)),
                    BlockKind.TABLE,
                    table_identifier=f"table-{table_number}",
                )
            )
        else:
            blocks.extend(_lines_to_blocks(text))
    return ExtractedDocument(blocks, metadata={"format": "html"})


def _csv(data: bytes, filename: str) -> ExtractedDocument:
    rows = list(csv.reader(io.StringIO(data.decode("utf-8-sig", errors="replace"))))
    if not rows:
        return ExtractedDocument([], metadata={"format": "csv"})
    header = rows[0]
    identifier = Path(filename).stem or "table"
    blocks = []
    for start in range(1, len(rows), 25):
        group = rows[start : start + 25]
        rendered = [" | ".join(header)]
        rendered.extend(
            f"Row {index}: " + " | ".join(row) for index, row in enumerate(group, start)
        )
        blocks.append(
            StructuralBlock(
                "\n".join(rendered),
                BlockKind.TABLE,
                table_identifier=identifier,
                row_range=f"{start}-{start + len(group) - 1}",
            )
        )
    return ExtractedDocument(blocks, metadata={"format": "csv", "headers": header})


def _source(data: bytes, extension: str, filename: str) -> ExtractedDocument:
    text = data.decode("utf-8", errors="replace")
    lines = text.splitlines()
    language = CODE_LANGUAGES[extension]
    blocks: list[StructuralBlock] = []
    if extension == ".py":
        try:
            tree = ast.parse(text)
            nodes = [
                node
                for node in tree.body
                if isinstance(
                    node,
                    (
                        ast.Import,
                        ast.ImportFrom,
                        ast.ClassDef,
                        ast.FunctionDef,
                        ast.AsyncFunctionDef,
                    ),
                )
            ]
            for node in nodes:
                end = getattr(node, "end_lineno", node.lineno)
                content = "\n".join(lines[node.lineno - 1 : end])
                blocks.append(
                    StructuralBlock(
                        content,
                        BlockKind.CODE,
                        line_range=f"{node.lineno}-{end}",
                        section=getattr(node, "name", "imports"),
                        metadata={"language": language, "source_path": Path(filename).name},
                    )
                )
        except SyntaxError:
            pass
    if not blocks:
        for start in range(0, len(lines), 80):
            blocks.append(
                StructuralBlock(
                    "\n".join(lines[start : start + 80]),
                    BlockKind.CODE,
                    line_range=f"{start + 1}-{min(len(lines), start + 80)}",
                    metadata={"language": language, "source_path": Path(filename).name},
                )
            )
    return ExtractedDocument(blocks, metadata={"format": "source_code", "language": language})


def _textual(data: bytes, *, markdown: bool) -> ExtractedDocument:
    return ExtractedDocument(
        _lines_to_blocks(data.decode("utf-8", errors="replace"), markdown=markdown)
    )


def _lines_to_blocks(
    text: str, *, page: int | None = None, markdown: bool = False
) -> list[StructuralBlock]:
    blocks: list[StructuralBlock] = []
    paragraph: list[str] = []
    in_fence = False
    fence: list[str] = []

    def flush() -> None:
        if paragraph:
            value = "\n".join(paragraph).strip()
            if value:
                blocks.append(StructuralBlock(value, BlockKind.PARAGRAPH, page=page))
            paragraph.clear()

    for raw in text.splitlines():
        line = raw.rstrip()
        if markdown and line.lstrip().startswith("```"):
            flush()
            fence.append(line)
            in_fence = not in_fence
            if not in_fence:
                blocks.append(StructuralBlock("\n".join(fence), BlockKind.CODE, page=page))
                fence.clear()
            continue
        if in_fence:
            fence.append(line)
            continue
        stripped = line.strip()
        if not stripped:
            flush()
            continue
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped) if markdown else None
        question = QUESTION.match(stripped)
        if heading_match:
            flush()
            title = heading_match.group(2).strip()
            blocks.append(
                StructuralBlock(
                    stripped,
                    BlockKind.HEADING,
                    page=page,
                    heading=title,
                    heading_level=len(heading_match.group(1)),
                    section=title,
                )
            )
        elif question:
            flush()
            blocks.append(
                StructuralBlock(
                    stripped, BlockKind.QUESTION, page=page, question_number=question.group(1)
                )
            )
        elif re.match(r"^\s*(?:[-*+]|\d+[.)])\s+\S", line):
            flush()
            blocks.append(StructuralBlock(stripped, BlockKind.LIST_ITEM, page=page))
        elif HEADING_VALUE.match(stripped):
            flush()
            heading = stripped.split(":", 1)[0]
            blocks.append(
                StructuralBlock(
                    stripped,
                    BlockKind.HEADING,
                    page=page,
                    heading=heading,
                    heading_level=2,
                    section=heading,
                )
            )
        else:
            paragraph.append(stripped)
    flush()
    if fence:
        blocks.append(StructuralBlock("\n".join(fence), BlockKind.CODE, page=page))
    return blocks
