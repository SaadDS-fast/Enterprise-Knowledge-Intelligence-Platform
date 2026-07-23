from __future__ import annotations

from datetime import UTC, datetime, timedelta
from enum import StrEnum
from hashlib import sha256
from hmac import compare_digest
from hmac import new as hmac_new
from typing import Any
from uuid import UUID

from pydantic import BaseModel, ConfigDict, Field, field_serializer, field_validator

from app.agents.schemas import AgentQueryResponse
from app.core.config import settings
from app.integrations.storage.keys import safe_object_name

PIPELINE_VERSION = "research-v1"


class ResearchState(StrEnum):
    PENDING = "PENDING"
    AUTHORIZING = "AUTHORIZING"
    PLANNING = "PLANNING"
    RETRIEVING = "RETRIEVING"
    RETRIEVAL_RETRY = "RETRIEVAL_RETRY"
    AGGREGATING_EVIDENCE = "AGGREGATING_EVIDENCE"
    VERIFYING_EVIDENCE = "VERIFYING_EVIDENCE"
    WRITING = "WRITING"
    VERIFYING_CITATIONS = "VERIFYING_CITATIONS"
    SAFETY_REVIEW = "SAFETY_REVIEW"
    EXPORTING = "EXPORTING"
    COMPLETED = "COMPLETED"
    FAILED = "FAILED"
    CANCEL_REQUESTED = "CANCEL_REQUESTED"
    CANCELLED = "CANCELLED"


ALLOWED_RESEARCH_TRANSITIONS: dict[ResearchState, frozenset[ResearchState]] = {
    ResearchState.PENDING: frozenset({ResearchState.AUTHORIZING, ResearchState.CANCEL_REQUESTED}),
    ResearchState.AUTHORIZING: frozenset({ResearchState.PLANNING, ResearchState.FAILED}),
    ResearchState.PLANNING: frozenset({ResearchState.RETRIEVING, ResearchState.FAILED}),
    ResearchState.RETRIEVING: frozenset(
        {ResearchState.RETRIEVAL_RETRY, ResearchState.AGGREGATING_EVIDENCE, ResearchState.FAILED}
    ),
    ResearchState.RETRIEVAL_RETRY: frozenset(
        {ResearchState.AGGREGATING_EVIDENCE, ResearchState.FAILED}
    ),
    ResearchState.AGGREGATING_EVIDENCE: frozenset(
        {ResearchState.VERIFYING_EVIDENCE, ResearchState.FAILED}
    ),
    ResearchState.VERIFYING_EVIDENCE: frozenset({ResearchState.WRITING, ResearchState.FAILED}),
    ResearchState.WRITING: frozenset({ResearchState.VERIFYING_CITATIONS, ResearchState.FAILED}),
    ResearchState.VERIFYING_CITATIONS: frozenset(
        {ResearchState.SAFETY_REVIEW, ResearchState.FAILED}
    ),
    ResearchState.SAFETY_REVIEW: frozenset({ResearchState.EXPORTING, ResearchState.FAILED}),
    ResearchState.EXPORTING: frozenset({ResearchState.COMPLETED, ResearchState.FAILED}),
    ResearchState.CANCEL_REQUESTED: frozenset({ResearchState.CANCELLED}),
    ResearchState.COMPLETED: frozenset(),
    ResearchState.FAILED: frozenset(),
    ResearchState.CANCELLED: frozenset(),
}


class ResearchFormat(StrEnum):
    MARKDOWN = "markdown"
    PDF = "pdf"
    DOCX = "docx"


class ResearchCreateRequest(BaseModel):
    question: str = Field(min_length=5, max_length=4000)
    document_ids: list[UUID] | None = None
    allow_external_sources: bool = False
    requested_formats: list[ResearchFormat] = Field(
        default_factory=lambda: [ResearchFormat.MARKDOWN]
    )
    max_depth_preset: str | None = Field(default="standard", max_length=40)
    idempotency_key: str | None = Field(default=None, min_length=8, max_length=160)

    @field_validator("requested_formats")
    @classmethod
    def validate_formats(cls, value: list[ResearchFormat]) -> list[ResearchFormat]:
        allowed = set(settings.agent_research_allowed_formats)
        requested = [item for item in value or [ResearchFormat.MARKDOWN]]
        if any(item.value not in allowed for item in requested):
            raise ValueError("Requested report format is not enabled")
        return sorted(set(requested), key=lambda item: item.value)


class ResearchArtifactRead(BaseModel):
    model_config = ConfigDict(from_attributes=True)

    id: UUID
    format: str
    filename: str
    mime_type: str
    checksum_sha256: str
    size_bytes: int
    created_at: datetime


class ResearchJobRead(BaseModel):
    model_config = ConfigDict(from_attributes=True)

    id: UUID
    workspace_id: UUID
    user_id: UUID
    agent_run_id: UUID | None = None
    request_id: str | None = None
    question: str
    status: str
    current_state: str
    stage: str
    progress_percent: int
    external_sources_allowed: bool
    requested_formats: list[str]
    source_count: int
    verified_citation_count: int
    artifact_refs: list[dict[str, Any]]
    error_code: str | None = None
    error_message: str | None = None
    result_json: dict[str, Any]
    created_at: datetime
    started_at: datetime | None = None
    completed_at: datetime | None = None
    cancelled_at: datetime | None = None
    updated_at: datetime

    @field_serializer("artifact_refs")
    def serialize_artifact_refs(self, value: list[dict[str, Any]]) -> list[dict[str, Any]]:
        allowed = {
            "artifact_id",
            "format",
            "filename",
            "mime_type",
            "checksum_sha256",
            "size_bytes",
            "signed_url_expires",
            "download_url",
        }
        return [{key: item[key] for key in allowed if key in item} for item in value]


class ResearchCreateResponse(BaseModel):
    job_id: UUID
    status: str
    current_state: str
    idempotent_replay: bool = False


class StructuredReport(BaseModel):
    title: str
    research_question: str
    executive_summary: str
    scope_and_methodology: str
    key_findings: list[str]
    detailed_analysis: str
    internal_evidence: list[dict[str, Any]]
    external_evidence: list[dict[str, Any]]
    conflicting_evidence: list[dict[str, Any]]
    information_gaps: list[str]
    limitations: list[str]
    conclusions: list[str]
    citations: list[dict[str, Any]]
    generation_metadata: dict[str, Any]


def can_transition(current: ResearchState, target: ResearchState) -> bool:
    return target in ALLOWED_RESEARCH_TRANSITIONS[current]


def validate_transition(current: str, target: ResearchState) -> None:
    state = ResearchState(current)
    if not can_transition(state, target):
        raise ValueError(f"Invalid research transition: {state.value} -> {target.value}")


def research_object_key(
    *, tenant_id: UUID, workspace_id: UUID, job_id: UUID, artifact_id: UUID, ext: str
) -> str:
    return f"reports/{tenant_id}/{workspace_id}/{job_id}/{artifact_id}/report.{ext}"


def scoped_idempotency_key(
    *, tenant_id: UUID, workspace_id: UUID, user_id: UUID, payload: ResearchCreateRequest
) -> str:
    provided = payload.idempotency_key or "none"
    documents = ",".join(str(item) for item in sorted(payload.document_ids or []))
    formats = ",".join(item.value for item in payload.requested_formats)
    raw = (
        f"{tenant_id}:{workspace_id}:{user_id}:{provided}:"
        f"{payload.question.strip().lower()}:{documents}:"
        f"{payload.allow_external_sources}:{formats}:{payload.max_depth_preset}"
    )
    return sha256(raw.encode("utf-8")).hexdigest()


def build_structured_report(response: AgentQueryResponse, question: str) -> StructuredReport:
    citations = response.citations
    outcome = response.outcome
    title = safe_object_name(question).replace("-", " ").title()[:120] or "Research Report"
    findings = [
        claim["claim_text"]
        for claim in response.claims
        if claim.get("verification_status") in {"SUPPORTED", "PARTIALLY_SUPPORTED"}
    ]
    if not findings and response.answer:
        findings = [response.answer]
    gaps = []
    if response.abstained or outcome in {"KNOWLEDGE_ABSENT", "INSUFFICIENT_EVIDENCE"}:
        gaps.append("The available authorized evidence did not fully answer the question.")
    if response.retrieval_diagnosis:
        gaps.append(
            f"Retrieval diagnosis: {response.retrieval_diagnosis.get('status', 'UNKNOWN')}."
        )
    return StructuredReport(
        title=title,
        research_question=question,
        executive_summary=response.answer or "No supported answer was produced.",
        scope_and_methodology=(
            "The report was generated asynchronously by the controlled agent using authorized "
            "workspace retrieval, multi-source evidence normalization, claim verification, "
            "citation validation, and safety review."
        ),
        key_findings=findings[:8],
        detailed_analysis=response.answer or "No detailed analysis was available.",
        internal_evidence=response.unified_evidence
        or [item.model_dump(mode="json") for item in response.internal_evidence],
        external_evidence=[item.model_dump(mode="json") for item in response.external_evidence],
        conflicting_evidence=response.conflicts,
        information_gaps=gaps,
        limitations=[
            "The report is limited to authorized internal documents and explicitly allowed "
            "external providers.",
            "Unsupported factual claims are removed or qualified.",
        ],
        conclusions=findings[:3] or ["No supported conclusion could be drawn."],
        citations=citations,
        generation_metadata={
            "pipeline_version": PIPELINE_VERSION,
            "agent_run_id": str(response.run_id),
            "outcome": outcome,
            "confidence_category": response.confidence_category,
            "fallback_used": response.fallback_used,
        },
    )


def render_markdown(report: StructuredReport, *, max_words: int | None = None) -> str:
    lines = [
        f"# {report.title}",
        "",
        "## Research Question",
        report.research_question,
        "",
        "## Executive Summary",
        report.executive_summary,
        "",
        "## Scope And Methodology",
        report.scope_and_methodology,
        "",
        "## Key Findings",
    ]
    lines.extend(f"- {item}" for item in report.key_findings)
    lines.extend(["", "## Detailed Analysis", report.detailed_analysis, ""])
    lines.append("## Conflicting Evidence")
    if report.conflicting_evidence:
        lines.extend(f"- {item.get('summary', item)}" for item in report.conflicting_evidence)
    else:
        lines.append("- None detected")
    lines.extend(["", "## Information Gaps"])
    if report.information_gaps:
        lines.extend(f"- {item}" for item in report.information_gaps)
    else:
        lines.append("- None")
    lines.extend(["", "## Limitations"])
    lines.extend(f"- {item}" for item in report.limitations)
    lines.extend(["", "## Conclusions"])
    lines.extend(f"- {item}" for item in report.conclusions)
    lines.extend(["", "## Citations"])
    if report.citations:
        lines.extend(
            (
                f"- [{item.get('citation_label') or item.get('external_source_label')}] "
                f"{item.get('document_title') or item.get('title')} "
                f"{item.get('canonical_url') or ''}"
            ).strip()
            for item in report.citations
        )
    else:
        lines.append("- None")
    lines.extend(
        [
            "",
            "## Generation Metadata",
            f"- Pipeline: {report.generation_metadata.get('pipeline_version')}",
            f"- Outcome: {report.generation_metadata.get('outcome')}",
            f"- Confidence: {report.generation_metadata.get('confidence_category')}",
        ]
    )
    content = "\n".join(lines).strip() + "\n"
    if max_words is None:
        max_words = settings.agent_research_max_report_words
    words = content.split()
    if len(words) > max_words:
        content = " ".join(words[:max_words])
    return content


def render_pdf(markdown: str) -> bytes:
    text = markdown.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    lines = text.splitlines()[:80]
    stream = "BT /F1 10 Tf 50 780 Td " + " T* ".join(f"({line[:100]}) Tj" for line in lines) + " ET"
    objects = [
        "1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj",
        "2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj",
        "3 0 obj << /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
        "/MediaBox [0 0 612 792] /Contents 5 0 R >> endobj",
        "4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj",
        f"5 0 obj << /Length {len(stream.encode())} >> stream\n{stream}\nendstream endobj",
    ]
    body = "%PDF-1.4\n" + "\n".join(objects) + "\ntrailer << /Root 1 0 R >>\n%%EOF\n"
    return body.encode("utf-8")


def render_docx(markdown: str) -> bytes:
    from io import BytesIO

    from docx import Document

    document = Document()
    for line in markdown.splitlines():
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            document.add_paragraph(line)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def sign_download_token(job_id: UUID, artifact_id: UUID, fmt: str) -> tuple[int, str]:
    expires = int(
        (
            datetime.now(UTC) + timedelta(seconds=settings.agent_research_signed_url_ttl_seconds)
        ).timestamp()
    )
    payload = f"{job_id}:{artifact_id}:{fmt}:{expires}"
    signature = hmac_new(
        settings.secret_key.get_secret_value().encode("utf-8"),
        payload.encode("utf-8"),
        "sha256",
    ).hexdigest()
    return expires, signature


def verify_download_token(
    job_id: UUID, artifact_id: UUID, fmt: str, expires: int, signature: str
) -> bool:
    if datetime.now(UTC).timestamp() > expires:
        return False
    payload = f"{job_id}:{artifact_id}:{fmt}:{expires}"
    expected = hmac_new(
        settings.secret_key.get_secret_value().encode("utf-8"),
        payload.encode("utf-8"),
        "sha256",
    ).hexdigest()
    return compare_digest(expected, signature)
