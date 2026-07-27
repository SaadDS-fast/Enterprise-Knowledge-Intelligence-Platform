import io

from docx import Document
from pypdf import PdfWriter

from app.ingestion.chunking_v3 import chunk_document
from app.ingestion.extractor import extract_document
from app.ingestion.quality import ExtractionQuality, assess_extraction
from app.ingestion.structure import BlockKind, ExtractedDocument
from app.ingestion.versions import LATEST_PIPELINE, is_current


def test_markdown_preserves_topic_equation_lists_and_fence() -> None:
    document = extract_document(
        ".md",
        b"""# Mathematics

Topic: Quadratic Equations

Definition:
A quadratic equation has the form ax\xc2\xb2 + bx + c = 0.

- Preserve symbols

```python
def solve():
    return 1
```
""",
    )
    chunks = chunk_document(document, chunk_size=300, overlap=20)
    assert "ax² + bx + c = 0" in document.text
    assert any(block.kind == BlockKind.CODE for block in document.blocks)
    assert any(chunk.metadata["heading"] == "Definition" for chunk in chunks)


def test_docx_preserves_heading_table_and_source_order() -> None:
    source = Document()
    source.add_heading("Travel Policy", level=1)
    source.add_paragraph("Domestic Meal Allowance: PKR 5,000 per day.")
    table = source.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Role"
    table.cell(0, 1).text = "Name"
    table.cell(1, 0).text = "Finance Director"
    table.cell(1, 1).text = "Ayesha Khan"
    payload = io.BytesIO()
    source.save(payload)

    document = extract_document(".docx", payload.getvalue())
    assert document.blocks[0].kind == BlockKind.HEADING
    assert any(block.kind == BlockKind.TABLE for block in document.blocks)
    assert document.text.index("Travel Policy") < document.text.index("Finance Director")


def test_csv_chunks_have_header_identifier_and_bounded_row_ranges() -> None:
    rows = ["name,value", *(f"item-{index},{index}" for index in range(60))]
    document = extract_document(".csv", "\n".join(rows).encode(), filename="metrics.csv")
    chunks = chunk_document(document, chunk_size=900, overlap=20)
    assert len(document.blocks) == 3
    assert document.blocks[0].table_identifier == "metrics"
    assert document.blocks[0].row_range == "1-25"
    assert all("name | value" in chunk.content for chunk in chunks)


def test_python_functions_are_not_merged() -> None:
    document = extract_document(
        ".py",
        b"import os\n\n\ndef first():\n    return 1\n\n\ndef second():\n    return 2\n",
        filename="example.py",
    )
    chunks = chunk_document(document)
    assert any(chunk.metadata["section"] == "first" for chunk in chunks)
    assert any(chunk.metadata["section"] == "second" for chunk in chunks)
    assert not any(
        "def first" in chunk.content and "def second" in chunk.content for chunk in chunks
    )


def test_quality_requires_ocr_for_empty_scanned_document() -> None:
    assessment = assess_extraction(
        ExtractedDocument([], page_count=4, empty_pages=4, scanned_likelihood=1.0)
    )
    assert assessment.status == ExtractionQuality.REQUIRES_OCR


def test_blank_image_only_pdf_requires_ocr_instead_of_parser_failure() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    payload = io.BytesIO()
    writer.write(payload)

    document = extract_document(".pdf", payload.getvalue())
    assert assess_extraction(document).status == ExtractionQuality.REQUIRES_OCR


def test_legacy_pipeline_metadata_is_not_current() -> None:
    assert not is_current({})
    assert is_current(LATEST_PIPELINE.as_dict())
